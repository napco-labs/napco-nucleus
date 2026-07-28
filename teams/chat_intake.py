"""File a chat dump a developer hands to the assistant into the central store.

Titu, 2026-07-28: "any developer give you all the chats he has, you will
transfer those in the central as well as these will be accountable while
identifying requirements."

The scheduled chat push (teams.push_chat) already ships every conversation the
assistant is part of, but it files everything under the assistant's own name
("napco-nucleus"), because that is whose Teams client it read. So a pile of
chats Rocky pastes in arrives at central looking like the assistant's own
chatter, and nothing ties it back to Rocky.

This module files it under the DEV who handed it over:

    <central>/<Dev>/<YYYY-MM-DD>/chat/handover_<HHMMSS>.docx

which is the same folder shape push_chat writes, so the requirement pipeline
picks it up with no change on the central side, and the requirements it yields
are attributable to the person who supplied them.

Deliberately NOT the requirement extractor: this only files the material. The
pipeline on central still decides what is a requirement.
"""
from __future__ import annotations

import datetime as dt
import os
import re
import socket
from pathlib import Path

_REPO = Path(__file__).parent.parent
LOCAL_OUT_DIR = _REPO / "data" / "teams" / "chat-handovers"

# A dump is a wall of conversation, not a question. Require real bulk AND
# several lines, so a long but ordinary message is never mistaken for one.
MIN_CHARS = 200
MIN_LINES = 4

# "[10:15] Rocky: ...", "Rocky, 10:15 AM", "10:15 AM Rocky" -- the shapes a
# pasted Teams/WhatsApp transcript actually takes.
_TS_LINE = re.compile(
    r"(\[\d{1,2}:\d{2}\s*(?:am|pm)?\]|"
    r"\b\d{1,2}:\d{2}\s*(?:am|pm)\b|"
    r"^\s*\S[^:\n]{0,40}:\s+\S)",
    re.I | re.M)


def looks_like_chat_dump(text: str) -> bool:
    """True when `text` reads as pasted conversation rather than a message."""
    t = (text or "").strip()
    if len(t) < MIN_CHARS:
        return False
    lines = [ln for ln in t.splitlines() if ln.strip()]
    if len(lines) < MIN_LINES:
        return False
    # at least a third of the lines carrying a timestamp or "Name:" prefix
    hits = sum(1 for ln in lines if _TS_LINE.search(ln))
    return hits >= max(3, len(lines) // 3)


def _central_chat_dir(dev: str) -> Path | None:
    raw = (os.environ.get("NUCLEUS_CENTRAL_PATH") or "").strip()
    if not raw:
        return None
    day = dt.date.today().strftime("%Y-%m-%d")
    return Path(raw) / dev / day / "chat"


def _build_docx(out_path: Path, dev: str, text: str, when: dt.datetime) -> None:
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    doc.add_heading("Teams chat handover", level=0)
    sub = doc.add_paragraph()
    sub.add_run(
        f"From: {dev}    "
        f"Received: {when:%Y-%m-%d %H:%M}    "
        f"Via: Napco Nucleus chat    "
        f"Host: {socket.gethostname()}"
    ).italic = True

    note = doc.add_paragraph()
    note_run = note.add_run(
        "Handed to the assistant directly by the developer named above. "
        "Treat these messages as that developer's own conversation history."
    )
    note_run.font.size = Pt(9)
    note_run.italic = True

    doc.add_heading(f"Chats supplied by {dev}", level=1)
    for line in (text or "").splitlines():
        if line.strip():
            doc.add_paragraph(line.rstrip())

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))


def file_handover(dev: str, text: str) -> tuple[bool, str]:
    """Write `text` to central under `dev`. Returns (ok, detail).

    Always keeps a local copy first, so a central share that is down loses
    nothing: the local docx can be copied up later.
    """
    dev = (dev or "").strip() or "unknown"
    when = dt.datetime.now()
    fname = f"handover_{when:%H%M%S}.docx"

    local = LOCAL_OUT_DIR / f"{when:%Y-%m-%d}" / dev / fname
    try:
        _build_docx(local, dev, text, when)
    except Exception as e:
        return False, f"could not write local copy: {str(e)[:120]}"

    dest_dir = _central_chat_dir(dev)
    if dest_dir is None:
        return False, "NUCLEUS_CENTRAL_PATH not set; kept local copy only"
    try:
        dest_dir.mkdir(parents=True, exist_ok=True)
        dest = dest_dir / fname
        import shutil
        shutil.copy2(str(local), str(dest))
        return True, str(dest)
    except Exception as e:
        return False, f"central copy failed ({str(e)[:120]}); kept {local}"
