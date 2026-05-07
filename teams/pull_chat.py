"""Pull messages from one Teams chat within a time window, write as a Word doc.

Resolves the chat by chat_number, conversation_id, OR substring match
against group title / participant names / past sender names. Supports
HH:MM and 12-hour ("3 PM", "3:30 PM") time formats.

Examples:
    py pull_chat.py --name "ContiHosting" --from 15:00 --to 17:00
    py pull_chat.py --name "Shafi" --from "3 PM" --to "5 PM"
    py pull_chat.py --number 123 --from 09:00 --to 18:00 --date 2026-05-06

Default output directory: NAPCO Nucleus's inbox/chat/ so the
requirement-management workflow picks the .docx up automatically.
Override with --out-dir or env CHAT_OUT_DIR.
"""
from __future__ import annotations

import argparse
import datetime as dt
import os
import re
import sqlite3
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt
from dotenv import load_dotenv

from teams import reader, store
from tools import _session_doc as session_doc

# Load NN's .env (CHAT_ALIASES + Teams config live there)
load_dotenv(Path(__file__).parent.parent / ".env", override=False)

DEFAULT_OUT_DIR = Path(
    os.environ.get("CHAT_OUT_DIR")
    or r"E:\Projects\NAPCO-Nucleus\data\requirements\inbox\chat"
)

# When --session is passed (default), the messages append to NN's
# pull-session doc INSTEAD of writing a standalone .docx in inbox/chat/.
# Pass --no-session to keep the legacy behavior.


def _load_aliases() -> dict[str, str]:
    """Return {alias_lowercase: identifier} from CHAT_ALIASES env var.

    Format:  CHAT_ALIASES=ContiHosting=123,InternalDevs=45,Conti=19:abc...skype
    Identifier can be a chat number or a full conversation_id.
    Aliases let users say --name "ContiHosting" even when the IndexedDB
    title is blank (common for legacy thread.skype chats).
    """
    raw = os.environ.get("CHAT_ALIASES", "").strip()
    if not raw:
        return {}
    out: dict[str, str] = {}
    for pair in raw.split(","):
        if "=" not in pair:
            continue
        name, ident = pair.split("=", 1)
        name = name.strip().lower()
        ident = ident.strip()
        if name and ident:
            out[name] = ident
    return out


def _resolve(arg: str):
    """Return (conversation_id, title, chat_number) or None."""
    if arg.isdigit():
        row = store.get_chat_by_number(int(arg))
        if row:
            return (row["conversation_id"],
                    row["title"] or "(untitled)",
                    row["chat_number"])
        return None

    if arg.startswith("19:"):
        row = store.get_chat_by_id(arg)
        if row:
            return (row["conversation_id"],
                    row["title"] or "(untitled)",
                    row["chat_number"])
        return (arg, "(unknown)", None)

    # Alias env-var lookup (CHAT_ALIASES=ContiHosting=123,...)
    aliases = _load_aliases()
    aliased = aliases.get(arg.lower())
    if aliased:
        # Recurse with the resolved identifier; show the alias in title fallback
        sub = _resolve(aliased)
        if sub:
            cid, title, number = sub
            if title in ("(untitled)", "(unknown)") or not title:
                title = arg  # use the alias as a friendly title
            return (cid, title, number)

    # Substring match across title, participants_json, then sender names
    needle = arg.lower()
    candidates: list[dict] = []

    for row in store.list_chat_registry(order="activity"):
        title = (row["title"] or "").lower()
        parts = (row["participants_json"] or "").lower()
        if needle in title or needle in parts:
            candidates.append({
                "conversation_id": row["conversation_id"],
                "title": row["title"] or f"(chat #{row['chat_number']})",
                "chat_number": row["chat_number"],
                "last_activity_ms": row["last_activity_ms"] or 0,
            })

    if not candidates:
        with sqlite3.connect(str(store.DB_PATH)) as conn:
            conn.row_factory = sqlite3.Row
            rows = conn.execute(
                """SELECT DISTINCT m.conversation_id,
                          r.title, r.chat_number, r.last_activity_ms
                   FROM messages m
                   LEFT JOIN chat_registry r
                          ON r.conversation_id = m.conversation_id
                   WHERE LOWER(m.sender_name) LIKE ?""",
                (f"%{needle}%",),
            ).fetchall()
            seen = set()
            for r in rows:
                cid = r["conversation_id"]
                if cid in seen:
                    continue
                seen.add(cid)
                candidates.append({
                    "conversation_id": cid,
                    "title": (r["title"] or
                              f"(chat with messages from {arg})"),
                    "chat_number": r["chat_number"],
                    "last_activity_ms": r["last_activity_ms"] or 0,
                })

    if not candidates:
        return None

    if len(candidates) > 1:
        print(f"Multiple chats matched {arg!r}; picking most recently active. "
              f"Other matches:")
        for c in sorted(candidates,
                        key=lambda r: r["last_activity_ms"], reverse=True)[1:6]:
            print(f"  #{c['chat_number']}  {c['title']}")

    best = max(candidates, key=lambda r: r["last_activity_ms"])
    return (best["conversation_id"], best["title"], best["chat_number"])


def _parse_time(s: str) -> dt.time:
    s = s.strip().upper().replace(".", "")
    for fmt in ("%H:%M", "%I:%M %p", "%I %p", "%I:%M%p", "%I%p"):
        try:
            return dt.datetime.strptime(s, fmt).time()
        except ValueError:
            continue
    raise ValueError(f"Unrecognized time {s!r}. Try '15:00' or '3:00 PM'.")


def _build_docx(out_path: Path, *, title: str, chat_number, cid: str,
                target_date: dt.date, from_t: dt.time, to_t: dt.time,
                msgs: list[dict], senders: list[str]) -> None:
    doc = Document()
    doc.add_heading(f"Teams chat: {title}", level=0)

    info = doc.add_paragraph()
    info.add_run("Date: ").bold = True
    info.add_run(f"{target_date}    ")
    info.add_run("Window: ").bold = True
    info.add_run(f"{from_t.strftime('%H:%M')} → {to_t.strftime('%H:%M')}\n")
    info.add_run("Chat: ").bold = True
    info.add_run(f"#{chat_number} {title}\n")
    info.add_run("Messages: ").bold = True
    info.add_run(f"{len(msgs)}    ")
    info.add_run("Participants: ").bold = True
    info.add_run(", ".join(senders) + "\n")
    cid_run = info.add_run(f"Conversation ID: {cid}")
    cid_run.font.size = Pt(8)

    doc.add_heading("Messages", level=1)
    for m in msgs:
        ts = dt.datetime.fromtimestamp((m.get("arrival_ms") or 0) / 1000)
        sender = m.get("sender_name") or "(unknown)"
        body = (m.get("body") or "").strip()
        if not body:
            continue
        p = doc.add_paragraph()
        p.add_run(f"[{ts:%H:%M}] {sender}: ").bold = True
        p.add_run(body)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))


def main() -> int:
    p = argparse.ArgumentParser(description=__doc__,
                                 formatter_class=argparse.RawDescriptionHelpFormatter)
    sel = p.add_mutually_exclusive_group(required=True)
    sel.add_argument("--name", help="Substring of group title, participant, or sender")
    sel.add_argument("--number", type=int, help="Chat number from registry")
    sel.add_argument("--id", help="Full conversation_id (19:...@thread.v2 / .skype)")
    p.add_argument("--last-minutes", type=int, default=None,
                   help="Pull the last N minutes (now - N min .. now). "
                        "Supersedes --from / --to / --date.")
    p.add_argument("--sender", default=None,
                   help="Filter messages to ones whose sender_name "
                        "contains this substring (case-insensitive). "
                        "E.g. --sender Salman.")
    p.add_argument("--from", dest="from_t", default="00:00",
                   help="Start time (HH:MM or '3 PM'). Default 00:00.")
    p.add_argument("--to", dest="to_t", default="23:59",
                   help="End time (HH:MM or '5 PM'). Default 23:59.")
    p.add_argument("--date", default=None,
                   help="Date YYYY-MM-DD. Default today.")
    p.add_argument("--out-dir", default=str(DEFAULT_OUT_DIR),
                   help="Where to write the standalone .docx (only used "
                        "with --no-session). Default: NN inbox/chat/.")
    p.add_argument("--no-session", action="store_true",
                   help="Write a standalone .docx in --out-dir instead of "
                        "appending to NN's pull-session doc.")
    args = p.parse_args()

    arg = args.id or (str(args.number) if args.number is not None else None) or args.name
    resolved = _resolve(arg)
    if not resolved:
        print(f"No chat matching {arg!r}.", file=sys.stderr)
        print("Run `py list_chats.py` to see available chats.", file=sys.stderr)
        return 2

    cid, title, number = resolved

    # Resolve absolute (start_dt, end_dt) from either --last-minutes
    # or --from/--to/--date. --last-minutes wins if both are given.
    if args.last_minutes is not None:
        if args.last_minutes <= 0:
            print("--last-minutes must be > 0", file=sys.stderr)
            return 1
        end_dt = dt.datetime.now()
        start_dt = end_dt - dt.timedelta(minutes=args.last_minutes)
        target_date = start_dt.date()
        from_t = start_dt.time()
        to_t = end_dt.time()
    else:
        target_date = (dt.datetime.strptime(args.date, "%Y-%m-%d").date()
                       if args.date else dt.date.today())
        try:
            from_t = _parse_time(args.from_t)
            to_t = _parse_time(args.to_t)
        except ValueError as e:
            print(f"Time parse error: {e}", file=sys.stderr)
            return 1
        start_dt = dt.datetime.combine(target_date, from_t)
        end_dt = dt.datetime.combine(target_date, to_t)

    from_ms = int(start_dt.timestamp() * 1000)
    to_ms = int(end_dt.timestamp() * 1000)

    print(f"Chat:    #{number}  {title}")
    if args.last_minutes is not None:
        print(f"Window:  last {args.last_minutes} min  "
              f"({start_dt:%Y-%m-%d %H:%M:%S} -> {end_dt:%H:%M:%S})")
    else:
        print(f"Window:  {target_date}  "
              f"{from_t.strftime('%H:%M')} -> {to_t.strftime('%H:%M')}")

    grouped = reader.read_messages_by_conversations({cid})
    msgs = grouped.get(cid, [])
    filtered = [m for m in msgs if from_ms <= (m.get("arrival_ms") or 0) <= to_ms]
    in_window = len(filtered)

    if args.sender:
        needle = args.sender.lower()
        filtered = [m for m in filtered
                    if needle in (m.get("sender_name") or "").lower()]
        print(f"Messages in window: {in_window} of {len(msgs)} total; "
              f"{len(filtered)} after --sender '{args.sender}' filter")
    else:
        print(f"Messages in window: {in_window} of {len(msgs)} total")

    if not filtered:
        print("No messages in the requested window. Nothing written.")
        return 0

    senders = sorted({m.get("sender_name") or "(unknown)" for m in filtered})

    if args.no_session:
        # Legacy path: standalone .docx in inbox/chat/
        safe = re.sub(r"[^A-Za-z0-9_-]+", "_",
                      (title or f"chat-{number}")).strip("_")[:40] or "chat"
        fname = (f"teams_{safe}_{target_date}_"
                 f"{from_t.strftime('%H%M')}-{to_t.strftime('%H%M')}.docx")
        out_path = Path(args.out_dir) / fname
        _build_docx(out_path, title=title, chat_number=number, cid=cid,
                    target_date=target_date, from_t=from_t, to_t=to_t,
                    msgs=filtered, senders=senders)
        print(f"\nWrote: {out_path.resolve()}")
        print(f"Participants: {', '.join(senders)}")
        return 0

    # Default path: append to NN's pull-session doc
    body_lines: list[str] = []
    for m in filtered:
        ts = dt.datetime.fromtimestamp((m.get("arrival_ms") or 0) / 1000)
        sender = m.get("sender_name") or "(unknown)"
        body = (m.get("body") or "").strip()
        if not body:
            continue
        body_lines.append(f"[{ts:%H:%M}] {sender}: {body}")

    headline = f"{title} (chat #{number})" if number else title
    if args.sender:
        headline += f"  · sender: {args.sender}"
    result = session_doc.append_section(
        source="TEAMS CHAT",
        headline=headline,
        metadata={
            "Date": str(target_date),
            "Window": f"{from_t.strftime('%H:%M')} -> {to_t.strftime('%H:%M')}",
            "Messages": str(len(filtered)),
            "Participants": ", ".join(senders),
        },
        body_paragraphs=body_lines,
    )

    print(f"\nAppended to session doc: {result['absolute_path']}")
    print(f"Section: {result['section']}")
    print(f"Lines added: {result['appended_paragraphs']}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
