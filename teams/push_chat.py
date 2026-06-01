"""Push the dev's recent Teams chat messages to the central store.

Runs on a 15-min Windows Task Scheduler timer per teammate. Each run:

  1. Reads every Teams chat in the local IndexedDB that has messages in
     the last N minutes (default 15).
  2. Consolidates them into ONE .docx (one section per chat with msgs).
  3. Writes the .docx locally to data/teams/chat-pushes/ for audit.
  4. If NUCLEUS_CENTRAL_PATH is set, copies the .docx to:
        <central>/<dev>/<YYYY-MM-DD>/chat/chat_<HHMM>-<HHMM>.docx
     so the agent host can pick up everyone's chat fragments in one place.

Differs from teams.pull_chat:
  - pull_chat appends to the per-dev session doc (the dev's local
    on-demand workflow). This script is for scheduled cross-dev sync.
  - Always writes a STANDALONE .docx per run; never touches session_doc.
  - All-chats only — no chat picking; no manual filters.

Usage
    python -m teams.push_chat                      # last 15 min
    python -m teams.push_chat --last-minutes 30
    python -m teams.push_chat --dry-run            # don't write anything

Exit codes
    0 — done (even if no messages found)
    2 — fatal config error (missing IndexedDB etc.)
"""
from __future__ import annotations

import argparse
import datetime as dt
import json
import os
import re
import shutil
import socket
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt
from dotenv import load_dotenv

_HERE = Path(__file__).parent.parent
load_dotenv(_HERE / ".env", override=True)

sys.path.insert(0, str(_HERE))

from teams import attachment_resolver, reader, store  # noqa: E402
from teams._exclude import excluded_conversation_ids  # noqa: E402


LOCAL_OUT_DIR = _HERE / "data" / "teams" / "chat-pushes"


def _dev_name() -> str:
    raw = (os.environ.get("NUCLEUS_DEV_NAME") or "").strip()
    if raw:
        return raw
    return (os.environ.get("USERNAME") or os.environ.get("USER")
            or socket.gethostname() or "unknown").strip()


def _central_chat_dir() -> Path | None:
    raw = (os.environ.get("NUCLEUS_CENTRAL_PATH") or "").strip()
    if not raw:
        return None
    day = dt.date.today().strftime("%Y-%m-%d")
    return Path(raw) / _dev_name() / day / "chat"


_UNSAFE_FNAME_CHARS = re.compile(r'[<>:"/\\|?*\x00-\x1f]')


def _safe_filename(name: str) -> str:
    """Strip path separators and Windows-forbidden chars from a filename."""
    base = Path(name).name or "attachment"
    return _UNSAFE_FNAME_CHARS.sub("_", base).strip(" .") or "attachment"


def _push_attachments(resolved: list[dict],
                      central_dir: Path) -> tuple[int, int]:
    """Copy each resolved attachment to <central_dir>/attachments/.
    Writes attachments.json manifest mapping URL -> filename.
    Returns (copied, skipped_existing)."""
    if not resolved:
        return (0, 0)
    att_dir = central_dir / "attachments"
    att_dir.mkdir(parents=True, exist_ok=True)
    manifest_path = att_dir / "manifest.json"
    manifest: dict[str, dict] = {}
    if manifest_path.exists():
        try:
            manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
        except Exception:
            manifest = {}

    copied = 0
    skipped = 0
    for att in resolved:
        local: Path = att["local_path"]
        safe_name = _safe_filename(att.get("name") or local.name)
        dst = att_dir / safe_name
        # Skip if already at central with matching size (the same file
        # may appear in multiple chat pushes during the same day).
        if dst.exists():
            try:
                if dst.stat().st_size == local.stat().st_size:
                    skipped += 1
                    manifest[att["url"]] = {
                        "name": att.get("name"),
                        "kind": att.get("kind"),
                        "size_bytes": att.get("size_bytes"),
                        "stored_as": safe_name,
                        "source_path": str(local),
                    }
                    continue
            except OSError:
                pass
            # name collision with different size — disambiguate
            stem, dot, ext = safe_name.rpartition(".")
            i = 1
            while dst.exists():
                cand = f"{stem} ({i}).{ext}" if dot else f"{safe_name} ({i})"
                dst = att_dir / cand
                i += 1
            safe_name = dst.name
        try:
            shutil.copy2(str(local), str(dst))
            copied += 1
            print(f"[push_chat]   attach: {safe_name} "
                  f"({local.stat().st_size / 1024:.1f} KB) <- {local}")
            manifest[att["url"]] = {
                "name": att.get("name"),
                "kind": att.get("kind"),
                "size_bytes": att.get("size_bytes"),
                "stored_as": safe_name,
                "source_path": str(local),
            }
        except Exception as e:
            print(f"[push_chat]   attach FAILED for {safe_name}: {e}",
                  file=sys.stderr)

    try:
        manifest_path.write_text(
            json.dumps(manifest, indent=2, ensure_ascii=False),
            encoding="utf-8")
    except Exception as e:
        print(f"[push_chat]   manifest write FAILED: {e}", file=sys.stderr)

    return copied, skipped


def _build_docx(out_path: Path,
                start_dt: dt.datetime, end_dt: dt.datetime,
                chat_blocks: list[dict]) -> None:
    """One Word doc, one heading per chat, message lines inside."""
    doc = Document()

    title = doc.add_heading("Teams chat push", level=0)
    sub = doc.add_paragraph()
    sub.add_run(
        f"Dev: {_dev_name()}    "
        f"Window: {start_dt:%Y-%m-%d %H:%M} -> {end_dt:%H:%M}    "
        f"Chats: {len(chat_blocks)}    "
        f"Host: {socket.gethostname()}"
    ).italic = True

    for blk in chat_blocks:
        title_str = (
            f"{blk['title']} (chat #{blk['chat_number']})"
            if blk.get("chat_number")
            else blk["title"]
        )
        doc.add_heading(title_str, level=1)
        info = doc.add_paragraph()
        info_run = info.add_run(
            f"Conversation: {blk['cid']}   "
            f"Msgs: {len(blk['msgs'])}   "
            f"Senders: {', '.join(blk['senders'])}"
        )
        info_run.font.size = Pt(9)
        info_run.italic = True

        for m in blk["msgs"]:
            ts = dt.datetime.fromtimestamp((m.get("arrival_ms") or 0) / 1000)
            sender = m.get("sender_name") or "(unknown)"
            body = (m.get("body") or "").strip()
            if not body:
                continue
            p = doc.add_paragraph()
            p.add_run(f"[{ts:%H:%M}] {sender}: ").bold = True
            p.add_run(body)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))


def _seed_registry_from_indexeddb() -> int:
    """Scan Teams IndexedDB and insert every group chat into chat_registry.

    Called automatically on first push_chat run when the registry is empty.
    Returns the number of chats seeded.
    """
    try:
        from ccl_chromium_reader import ccl_chromium_indexeddb
    except ImportError:
        print("[push_chat] ccl_chromium_reader not installed — cannot seed",
              file=sys.stderr)
        return 0

    leveldb = reader.LEVELDB_PATH
    if not leveldb.exists():
        return 0

    try:
        db = ccl_chromium_indexeddb.WrappedIndexDB(str(leveldb))
    except Exception as e:
        print(f"[push_chat] IndexedDB open failed: {e}", file=sys.stderr)
        return 0

    # Load participant display names
    profiles: dict[str, str] = {}
    for info in db.database_ids:
        if info.name and info.name.startswith("Teams:profiles:"):
            try:
                for rec in db[info.dbid_no]["profiles"].iterate_records():
                    v = rec.value
                    if isinstance(v, dict):
                        mri = v.get("mri") or v.get("id")
                        name = v.get("displayName") or v.get("givenName") or ""
                        if mri and name:
                            profiles[mri] = name
            except Exception:
                pass
            break

    # Collect all group chats from conversation-manager
    seeded = 0
    for info in db.database_ids:
        if info.name and info.name.startswith("Teams:conversation-manager:"):
            try:
                for rec in db[info.dbid_no]["conversations"].iterate_records():
                    v = rec.value
                    if not isinstance(v, dict):
                        continue
                    cid = v.get("id") or ""
                    if "thread.v2" not in cid:
                        continue
                    last_ms = v.get("lastMessageTimeUtc")
                    title = v.get("cachedDeduplicationKey") or v.get("title") or None
                    store.upsert_chat(
                        conversation_id=cid,
                        title=title,
                        fmt="group",
                        last_activity_ms=last_ms,
                        participants=[],
                        msg_count=0,
                    )
                    seeded += 1
            except Exception as e:
                print(f"[push_chat] seed error: {e}", file=sys.stderr)
            break

    return seeded


def main() -> int:
    ap = argparse.ArgumentParser(
        description=__doc__,
        formatter_class=argparse.RawDescriptionHelpFormatter)
    ap.add_argument("--last-minutes", type=int, default=15,
                    help="Window size. Default: 15.")
    ap.add_argument("--no-attachments", dest="push_attachments",
                    action="store_false",
                    help="Skip resolving + pushing chat-attached files.")
    ap.set_defaults(push_attachments=True)
    ap.add_argument("--dry-run", action="store_true",
                    help="Read + summarize but write nothing.")
    args = ap.parse_args()

    if args.last_minutes <= 0:
        print("--last-minutes must be > 0", file=sys.stderr)
        return 1

    end_dt = dt.datetime.now()
    start_dt = end_dt - dt.timedelta(minutes=args.last_minutes)
    from_ms = int(start_dt.timestamp() * 1000)
    to_ms = int(end_dt.timestamp() * 1000)

    print(f"[push_chat] dev={_dev_name()} host={socket.gethostname()}")
    print(f"[push_chat] window: {start_dt:%Y-%m-%d %H:%M} -> {end_dt:%H:%M}")

    rows = store.list_chat_registry(order="activity")
    if not rows:
        # Registry empty on first run — auto-seed from Teams IndexedDB so
        # the scheduled task works without any manual setup step.
        print("[push_chat] registry empty — auto-seeding from Teams IndexedDB...")
        seeded = _seed_registry_from_indexeddb()
        if seeded == 0:
            print("[push_chat] no chats found in IndexedDB. "
                  "Is Teams installed and has the user signed in?", file=sys.stderr)
            return 2
        print(f"[push_chat] seeded {seeded} chat(s) into registry")
        rows = store.list_chat_registry(order="activity")
    # store.list_chat_registry returns sqlite3.Row; flatten to plain dicts
    # so .get() works downstream (Row only supports __getitem__).
    chat_lookup: dict[str, dict] = {
        r["conversation_id"]: {
            "title": r["title"],
            "chat_number": r["chat_number"],
        }
        for r in rows
    }

    excluded = excluded_conversation_ids()
    if excluded:
        before = len(chat_lookup)
        dropped = [
            (cid, info) for cid, info in chat_lookup.items() if cid in excluded
        ]
        chat_lookup = {
            cid: info for cid, info in chat_lookup.items() if cid not in excluded
        }
        if dropped:
            print(f"[push_chat] excluding {len(dropped)} chat(s) via "
                  f"NUCLEUS_EXCLUDE_CHATS ({before} -> {len(chat_lookup)}):")
            for cid, info in dropped:
                label = info.get("title") or f"chat #{info.get('chat_number')}"
                print(f"   - {label}  ({cid})")

    cids = set(chat_lookup.keys())
    grouped = reader.read_messages_by_conversations(cids, since_ms=from_ms)

    chat_blocks: list[dict] = []
    total_msgs = 0
    for cid, msgs in grouped.items():
        if cid in excluded:
            # Defensive: registry-based filter above should catch this,
            # but reader.read_messages_by_conversations may surface
            # messages from a freshly seen group not yet in the registry.
            continue
        windowed = [m for m in msgs
                    if from_ms <= (m.get("arrival_ms") or 0) <= to_ms
                    and (m.get("body") or "").strip()]
        if not windowed:
            continue
        senders = sorted({m.get("sender_name") or "(unknown)" for m in windowed})
        info = chat_lookup.get(cid, {})
        chat_no = info.get("chat_number")
        title = info.get("title") or (
            f"(chat #{chat_no})" if chat_no else "(untitled)")
        chat_blocks.append({
            "cid": cid,
            "title": title,
            "chat_number": chat_no,
            "msgs": windowed,
            "senders": senders,
        })
        total_msgs += len(windowed)

    if not chat_blocks:
        print(f"[push_chat] no messages in window across "
              f"{len(cids)} chats. Nothing to push.")
        return 0

    print(f"[push_chat] {len(chat_blocks)} chat(s) with {total_msgs} msg(s) total")

    name = (
        f"chat_{start_dt:%Y-%m-%d}_"
        f"{start_dt:%H%M}-{end_dt:%H%M}.docx"
    )

    # Resolve Teams chat attachments to local files. We only ship the
    # ones the dev has actually downloaded; the rest stay as URL-only
    # references inside the .docx body.
    resolved: list[dict] = []
    if args.push_attachments:
        all_atts = attachment_resolver.resolve_all(chat_blocks)
        resolved = [a for a in all_atts if a["resolved"]]
        n_total = len(all_atts)
        n_found = len(resolved)
        if n_total:
            print(f"[push_chat] attachments seen: {n_total}, "
                  f"locally resolved: {n_found}")
            for a in all_atts:
                if a["resolved"]:
                    print(f"   + {a['name']}  ({a['kind']}, "
                          f"{a.get('size_bytes', 0)} B)  <- {a['local_path']}")
                else:
                    print(f"   - {a['name']}  ({a['kind']})  "
                          f"skipped: {a['reason']}")

    if args.dry_run:
        print(f"[push_chat] dry-run: would write {name}")
        for blk in chat_blocks:
            print(f"   + {blk['title']}  ({len(blk['msgs'])} msg)")
        return 0

    local_path = LOCAL_OUT_DIR / name
    _build_docx(local_path, start_dt, end_dt, chat_blocks)
    size_kb = local_path.stat().st_size / 1024
    print(f"[push_chat] local: {local_path}  ({size_kb:.1f} KB)")

    central_dir = _central_chat_dir()
    if central_dir is None:
        print("[push_chat] central upload: skipped "
              "(NUCLEUS_CENTRAL_PATH not set)")
        return 0
    from teams._central import ensure_smb_auth
    ensure_smb_auth(os.environ.get("NUCLEUS_CENTRAL_PATH", ""))
    try:
        central_dir.mkdir(parents=True, exist_ok=True)
        dst = central_dir / name
        shutil.copy2(str(local_path), str(dst))
        print(f"[push_chat] -> {dst}")
        print("[push_chat] central upload: OK")
    except Exception as e:
        print(f"[push_chat] central upload FAILED: {e}", file=sys.stderr)
        print(f"[push_chat] local copy preserved at {local_path}",
              file=sys.stderr)
        # Return non-zero so the Task Scheduler tick records the
        # failure. Previously we returned 0, making a stuck Samba share
        # invisible to the scheduler -- chat docs would pile up locally
        # while the supervisor reported success every tick.
        return 1

    if resolved:
        try:
            copied, skipped = _push_attachments(resolved, central_dir)
            print(f"[push_chat] attachments -> central: "
                  f"{copied} copied, {skipped} already present")
        except Exception as e:
            print(f"[push_chat] attachment push FAILED: {e}",
                  file=sys.stderr)
    return 0


if __name__ == "__main__":
    sys.exit(main())
