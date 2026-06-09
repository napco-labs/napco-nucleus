"""
NAPCO Nucleus — Word document writers for requirement collection.

Two MCP tools:

    write_aggregation_docx     Raw text per channel (email/chat/meeting/document)
                               for traceability — what came in.
    write_verification_docx    Identified requirements as 2-5 paragraph summaries
                               for client verification — what we think they said.

Both write into data/requirements/ with a date stamp. The agent passes
already-prepared structured input — no LLM reasoning happens here.
"""
from __future__ import annotations

import json
import logging
from datetime import datetime
from pathlib import Path

from claude_agent_sdk import tool

import memory

logger = logging.getLogger(__name__)

_HERE = Path(__file__).parent.parent
_OUT_DIR = _HERE / "data" / "requirements"


def _text(payload) -> dict:
    return {"content": [{"type": "text", "text": json.dumps(payload, default=str)}]}


def _today_stamp() -> str:
    return datetime.now().strftime("%Y-%m-%d")


# ─── write_aggregation_docx ─────────────────────────────────────────

@tool(
    "write_aggregation_docx",
    "Write the raw aggregation Word document — one section per channel "
    "(email / chat / meeting / document) listing every collected source "
    "with filename + full text. This is the traceability artifact the "
    "human reviewer reads alongside the verification doc. Input `sources` "
    "is a list of dicts with keys: channel (one of email/chat/meeting/"
    "document), filename (str), content (str — the full extracted text). "
    "Output path defaults to data/requirements/aggregation_<YYYY-MM-DD>.docx. "
    "Returns {path, channel_counts, total_sources}.",
    {"sources": list, "output_path": str},
)
async def write_aggregation_docx_tool(args):
    from docx import Document  # lazy
    from docx.shared import Pt

    sources = args.get("sources") or []
    if not isinstance(sources, list):
        return _text({"error": "sources must be a list"})

    out_path = args.get("output_path")
    if not out_path:
        _OUT_DIR.mkdir(parents=True, exist_ok=True)
        out_path = str(_OUT_DIR / f"aggregation_{_today_stamp()}.docx")

    doc = Document()

    title = doc.add_heading("Requirement Collection — Raw Aggregation", level=0)
    sub = doc.add_paragraph()
    sub.add_run(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}").italic = True
    doc.add_paragraph(
        "This document collects every raw input the agent ingested across "
        "Email, Teams chat, and Teams call audio (transcribed). Use it to "
        "verify nothing was missed before reviewing the Verification doc."
    )
    doc.add_paragraph()

    channel_order = ["email", "chat", "meeting", "document"]
    channel_titles = {
        "email": "Email",
        "chat": "Teams Chat",
        "meeting": "Teams Call (Transcribed)",
        "document": "Documents / Attachments",
    }
    bucket: dict[str, list[dict]] = {c: [] for c in channel_order}
    for s in sources:
        if not isinstance(s, dict):
            continue
        ch = (s.get("channel") or "").strip().lower()
        if ch not in bucket:
            ch = "document"
        bucket[ch].append(s)

    counts: dict[str, int] = {}
    for ch in channel_order:
        items = bucket[ch]
        counts[ch] = len(items)
        if not items:
            continue
        doc.add_heading(f"{channel_titles[ch]} ({len(items)})", level=1)
        for s in items:
            fname = (s.get("filename") or "(unnamed)").strip()
            content = (s.get("content") or "").strip()
            doc.add_heading(fname, level=2)
            if not content:
                p = doc.add_paragraph()
                p.add_run("(empty)").italic = True
                continue
            for line in content.splitlines():
                p = doc.add_paragraph(line)
                for run in p.runs:
                    run.font.size = Pt(10)
        doc.add_paragraph()

    Path(out_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)

    memory.log_activity(
        task_name="requirement-collection:write_aggregation",
        result=f"sources={len(sources)} path={Path(out_path).name}",
        technical_details={"channel_counts": counts, "path": out_path},
    )

    return _text({
        "path": out_path,
        "channel_counts": counts,
        "total_sources": len(sources),
    })


# ─── write_verification_docx ────────────────────────────────────────

@tool(
    "write_verification_docx",
    "Write the client-facing verification Word document as a flat numbered "
    "list. Each requirement is one numbered paragraph: '1. <title> - "
    "<one-paragraph summary>'. Tone: neutral developer English, no jargon, "
    "no marketing voice. Input `requirements` is a list of dicts with keys: "
    "title (str - short, <80 chars), summary (str - exactly ONE paragraph, "
    "no blank lines), source_refs (optional list of Source IDs from the "
    "pull-session metadata, e.g. 'chat/123/...', 'email/.../...'), "
    "confidence (optional float 0.0-1.0 - the LLM's certainty this is a "
    "real client requirement), rationale (optional str - one short "
    "sentence on why this counts as a requirement, not noise), "
    "priority (optional 'P0'|'P1'|'P2'|'P3' - urgency tag), "
    "severity (optional 'S1'|'S2'|'S3' - blast radius tag), "
    "estimate_hours (optional int - effort for this single workable task, "
    "targeting ~4h; rendered as a '~Nh' tag in the [priority/severity] "
    "bracket), "
    "conflicts_with (optional list of Source IDs or open-item ids the "
    "requirement appears to contradict — rendered as an amber WARNING "
    "line so the reviewer notices), "
    "time_ranges (optional list of {source_id, start, end} dicts where "
    "start/end are 'HH:MM:SS' clock times within a MEETING source. "
    "Lets the reviewer pull a short audio snippet via "
    "tools/audio_snippet.py to spot-check call-derived requirements). "
    "Output filename format is 'Requirements Verification <YYYY-MM-DD>.docx' "
    "in data/requirements/. Returns {path, requirement_count}.",
    {"requirements": list, "output_path": str},
)
async def write_verification_docx_tool(args):
    from docx import Document  # lazy
    from docx.shared import Pt, RGBColor

    reqs = args.get("requirements") or []
    if not isinstance(reqs, list) or not reqs:
        return _text({"error": "requirements must be a non-empty list"})

    out_path = args.get("output_path")
    if not out_path:
        _OUT_DIR.mkdir(parents=True, exist_ok=True)
        out_path = str(_OUT_DIR / f"Requirements Verification {_today_stamp()}.docx")

    doc = Document()

    doc.add_heading("Requirements Verification", level=0)
    sub = doc.add_paragraph()
    sub.add_run(f"Date: {datetime.now().strftime('%Y-%m-%d')}").italic = True

    intro = doc.add_paragraph()
    intro.add_run(
        "Below are the requirements identified from your recent communications "
        "(email, Teams messages, and call discussions). Please review each "
        "item and reply to confirm the interpretation, or send corrections. "
        "Each item will be tracked separately once you confirm."
    )
    doc.add_paragraph()

    _GREY = RGBColor(0x55, 0x60, 0x70)
    _AMBER = RGBColor(0xB8, 0x7C, 0x1F)

    def _confidence_color(c: float) -> RGBColor:
        # Low-confidence items get a soft amber accent so the reviewer
        # eyeballs them more carefully before the email goes out.
        return _AMBER if c < 0.75 else _GREY

    _PRIORITY_VALID = {"P0", "P1", "P2", "P3"}
    _SEVERITY_VALID = {"S1", "S2", "S3"}

    for i, r in enumerate(reqs, 1):
        if not isinstance(r, dict):
            continue
        title = (r.get("title") or "").strip() or f"Requirement {i}"
        summary = (r.get("summary") or "").strip()
        rationale = (r.get("rationale") or "").strip()
        source_refs = r.get("source_refs") or []
        confidence = r.get("confidence")
        try:
            conf_val = float(confidence) if confidence is not None else None
        except (TypeError, ValueError):
            conf_val = None
        priority = (r.get("priority") or "").strip().upper()
        if priority not in _PRIORITY_VALID:
            priority = ""
        severity = (r.get("severity") or "").strip().upper()
        if severity not in _SEVERITY_VALID:
            severity = ""
        estimate = r.get("estimate_hours")
        try:
            est_val = int(round(float(estimate))) if estimate is not None else None
        except (TypeError, ValueError):
            est_val = None
        if est_val is not None and est_val <= 0:
            est_val = None
        conflicts = r.get("conflicts_with") or []
        if not isinstance(conflicts, list):
            conflicts = []
        time_ranges = r.get("time_ranges") or []
        if not isinstance(time_ranges, list):
            time_ranges = []

        # Plain label (Titu's spec 2026-06-09): "Requirement#1: <title> -
        # <summary>". No priority/severity/effort tag in the client doc —
        # those stay in the JSON sidecar for internal / OpenProject use only.
        p = doc.add_paragraph()
        p.add_run(f"Requirement#{i}: ").bold = True
        p.add_run(title).bold = True
        if summary:
            collapsed = " ".join(line.strip() for line in summary.splitlines() if line.strip())
            p.add_run(f" - {collapsed}")
        else:
            p.add_run(" - ")
            p.add_run("(no summary supplied)").italic = True

        # Conflict warning line — separate paragraph in bold amber so
        # the reviewer can't miss it.
        if conflicts:
            warn_p = doc.add_paragraph()
            warn_run = warn_p.add_run(
                "⚠ Possible conflict with: " +
                ", ".join(str(c) for c in conflicts)
            )
            warn_run.bold = True
            warn_run.italic = True
            warn_run.font.size = Pt(9)
            warn_run.font.color.rgb = _AMBER

        # Audio time-range hints for MEETING-derived requirements.
        # Reviewer pulls a snippet via tools/audio_snippet.py.
        if time_ranges:
            tr_p = doc.add_paragraph()
            valid = []
            for tr in time_ranges:
                if not isinstance(tr, dict):
                    continue
                sid = (tr.get("source_id") or "").strip()
                st = (tr.get("start") or "").strip()
                en = (tr.get("end") or "").strip()
                if sid and st and en:
                    valid.append(f"{sid} @ {st}-{en}")
            if valid:
                tr_run = tr_p.add_run("🎧 Audio: " + " | ".join(valid))
                tr_run.italic = True
                tr_run.font.size = Pt(9)
                tr_run.font.color.rgb = _GREY

        # Citation + confidence + rationale on one grey 9pt line (or
        # split across lines if it gets long).
        bits: list[tuple[str, RGBColor]] = []
        if isinstance(source_refs, list) and source_refs:
            bits.append(
                ("Sources: " + ", ".join(str(s) for s in source_refs), _GREY))
        if conf_val is not None:
            color = _confidence_color(conf_val)
            label = f"Confidence: {conf_val:.2f}"
            if conf_val < 0.75:
                label += " (review)"
            bits.append((label, color))
        if rationale:
            bits.append((f"Why: {rationale}", _GREY))

        if bits:
            ref_p = doc.add_paragraph()
            for j, (text, color) in enumerate(bits):
                if j > 0:
                    sep = ref_p.add_run("   ·   ")
                    sep.italic = True
                    sep.font.size = Pt(9)
                    sep.font.color.rgb = _GREY
                run = ref_p.add_run(text)
                run.italic = True
                run.font.size = Pt(9)
                run.font.color.rgb = color

    doc.add_paragraph()
    closing = doc.add_paragraph()
    closing.add_run(
        "Please reply to this email confirming the above interpretation is "
        "correct, or send any corrections inline. Once confirmed, each item "
        "will be filed for development."
    )

    Path(out_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)

    # JSON sidecar — the eval harness reads this to score predicted vs
    # expected without having to parse the .docx. Same basename as the
    # .docx, .json extension. Always written so the file is available
    # for ad-hoc inspection too.
    sidecar_path = str(Path(out_path).with_suffix(".json"))
    try:
        sidecar = {
            "generated_at": datetime.now().isoformat(timespec="seconds"),
            "docx_path": out_path,
            "requirement_count": len(reqs),
            "requirements": [
                {
                    "title": r.get("title"),
                    "summary": r.get("summary"),
                    "source_refs": r.get("source_refs") or [],
                    "confidence": r.get("confidence"),
                    "rationale": r.get("rationale"),
                    "priority": (r.get("priority") or "").strip().upper() or None,
                    "severity": (r.get("severity") or "").strip().upper() or None,
                    "estimate_hours": r.get("estimate_hours"),
                    "conflicts_with": r.get("conflicts_with") or [],
                    "time_ranges": [
                        {"source_id": (tr.get("source_id") or "").strip(),
                         "start": (tr.get("start") or "").strip(),
                         "end": (tr.get("end") or "").strip()}
                        for tr in (r.get("time_ranges") or [])
                        if isinstance(tr, dict)
                    ],
                }
                for r in reqs if isinstance(r, dict)
            ],
        }
        Path(sidecar_path).write_text(
            json.dumps(sidecar, indent=2, ensure_ascii=False, default=str),
            encoding="utf-8",
        )
    except Exception as e:
        logger.warning("sidecar write failed for %s: %s", sidecar_path, e)
        sidecar_path = None

    # Telemetry for the eval harness (Phase 2) to read later — keeps
    # per-requirement confidence + citation count on the side.
    titles = []
    confidences = []
    citation_counts = []
    for r in reqs:
        if not isinstance(r, dict):
            continue
        titles.append(r.get("title"))
        c = r.get("confidence")
        try:
            confidences.append(float(c) if c is not None else None)
        except (TypeError, ValueError):
            confidences.append(None)
        srcs = r.get("source_refs") or []
        citation_counts.append(len(srcs) if isinstance(srcs, list) else 0)

    valid_conf = [c for c in confidences if c is not None]
    mean_conf = sum(valid_conf) / len(valid_conf) if valid_conf else None
    low_conf = sum(1 for c in valid_conf if c < 0.75)

    memory.log_activity(
        task_name="requirement-collection:write_verification",
        result=(f"requirements={len(reqs)} "
                f"mean_conf={mean_conf:.2f} " if mean_conf is not None else
                f"requirements={len(reqs)} ") +
               f"low_conf={low_conf} path={Path(out_path).name}",
        technical_details={
            "path": out_path,
            "titles": titles,
            "confidences": confidences,
            "citation_counts": citation_counts,
            "mean_confidence": mean_conf,
            "low_confidence_count": low_conf,
        },
    )

    return _text({
        "path": out_path,
        "sidecar_path": sidecar_path,
        "requirement_count": len(reqs),
        "mean_confidence": mean_conf,
        "low_confidence_count": low_conf,
    })


TOOLS = [
    write_aggregation_docx_tool,
    write_verification_docx_tool,
]

TOOL_NAMES = [
    "write_aggregation_docx",
    "write_verification_docx",
]
