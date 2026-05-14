"""
Pull-session Word document — single consolidated artifact that every
on-demand pull command (Teams chat, email, Drive file) appends to.

A "session" is a logical grouping of pulls the user wants identified
together. The session resets ONLY when the user explicitly commands
"start a new session" — pulls otherwise keep appending to the same doc
across days, terminals, and tool calls.

Layout:
    data/requirements/sessions/current.docx              <- the live session doc
    data/requirements/sessions/.current_meta.json        <- start time + label
    data/requirements/sessions/archive/<timestamp>.docx  <- previous sessions

Format inside current.docx:
    Title  : "Pull Session — started <timestamp>"
    Heading1 per pull section, e.g.:
        "TEAMS CHAT — ContiHosting (chat #123)"
        "EMAIL — from titucse@gmail.com (subject: 'budget')"
        "DRIVE — file requirements_v2.pdf"
    Each section: small metadata block, then the raw content as paragraphs.

Used by:
    - TRW: pull_chat.py
    - NN:  pull_email.py, pull_drive.py
    - NN:  the verify-session prompt (reads only this file)
"""
from __future__ import annotations

import hashlib
import json
import re
import shutil
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor

# ── Paths ─────────────────────────────────────────────────────────────
_NN_ROOT = Path(__file__).parent.parent
SESSIONS_DIR = _NN_ROOT / "data" / "requirements" / "sessions"
SESSION_PATH = SESSIONS_DIR / "current.docx"
META_PATH = SESSIONS_DIR / ".current_meta.json"
ARCHIVE_DIR = SESSIONS_DIR / "archive"

# ── Style constants ───────────────────────────────────────────────────
_NAVY = RGBColor(0x1F, 0x3A, 0x5F)
_GREY = RGBColor(0x55, 0x60, 0x70)


def _now_iso() -> str:
    return datetime.now().strftime("%Y-%m-%d %H:%M:%S")


def _stamp() -> str:
    return datetime.now().strftime("%Y%m%d-%H%M%S")


def _load_meta() -> dict:
    if not META_PATH.exists():
        return {}
    try:
        return json.loads(META_PATH.read_text(encoding="utf-8"))
    except Exception:
        return {}


def _save_meta(meta: dict) -> None:
    META_PATH.parent.mkdir(parents=True, exist_ok=True)
    META_PATH.write_text(json.dumps(meta, indent=2, default=str), encoding="utf-8")


def _create_empty(label: str | None = None) -> dict:
    """Initialize a fresh session doc and meta. Returns the meta dict."""
    SESSIONS_DIR.mkdir(parents=True, exist_ok=True)
    started = _now_iso()
    meta = {
        "started_at": started,
        "label": label or "",
        "session_path": str(SESSION_PATH.relative_to(_NN_ROOT).as_posix()),
    }

    doc = Document()
    title = doc.add_heading("Pull Session", level=0)
    sub = doc.add_paragraph()
    r = sub.add_run(f"Started {started}")
    r.italic = True
    r.font.color.rgb = _GREY
    if label:
        sub.add_run(f"   |   Label: {label}").italic = True
    doc.add_paragraph()  # spacer
    doc.save(str(SESSION_PATH))

    _save_meta(meta)
    return meta


def get_or_create() -> Path:
    """Return the current session doc path, creating it if absent."""
    if not SESSION_PATH.exists():
        _create_empty()
    return SESSION_PATH


def reset(label: str | None = None) -> dict:
    """Archive the current session (if any) and start a fresh one.

    Returns: {archived_to: str|None, new_started_at: str, new_label: str}
    """
    archived_to = None
    if SESSION_PATH.exists():
        ARCHIVE_DIR.mkdir(parents=True, exist_ok=True)
        old_meta = _load_meta()
        old_label = (old_meta.get("label") or "").strip()
        suffix = f"_{old_label}" if old_label else ""
        # Sanitize label for filename
        suffix = "".join(c for c in suffix if c.isalnum() or c in "-_")[:32]
        target = ARCHIVE_DIR / f"session_{_stamp()}{suffix}.docx"
        shutil.move(str(SESSION_PATH), str(target))
        archived_to = str(target.relative_to(_NN_ROOT).as_posix())
        try:
            META_PATH.unlink()
        except Exception:
            pass

    meta = _create_empty(label=label)
    return {
        "archived_to": archived_to,
        "new_started_at": meta["started_at"],
        "new_label": meta.get("label", ""),
        "session_path": meta["session_path"],
    }


_SLUG_RE = re.compile(r"[^a-zA-Z0-9._-]+")


def _slugify(s: str, max_len: int = 60) -> str:
    """Stable, URL-ish, machine-friendly slug for use inside a source_id."""
    s = _SLUG_RE.sub("-", (s or "").strip()).strip("-")
    return s[:max_len] or "x"


def _derive_source_id(source: str, headline: str) -> str:
    """Auto-derive a stable source_id when the caller didn't supply one.

    Format: <source_prefix>/<headline-slug>/<8-char-hash>. The hash keeps
    the ID stable across runs but unique even when two sections share
    similar headlines.
    """
    prefix_map = {
        "TEAMS CHAT": "chat",
        "TEAMS ATTACHMENT": "chat-attach",
        "EMAIL": "email",
        "MEETING": "call",
        "DRIVE": "drive",
    }
    prefix = prefix_map.get(source.upper(), _slugify(source).lower())
    slug = _slugify(headline)
    h = hashlib.sha1(f"{source}|{headline}".encode("utf-8")).hexdigest()[:8]
    return f"{prefix}/{slug}/{h}"


def append_section(
    *,
    source: str,                # "TEAMS CHAT" / "EMAIL" / "DRIVE" / "MEETING"
    headline: str,              # e.g. 'ContiHosting (chat #123)' or 'from titucse@gmail.com'
    metadata: dict[str, str],   # key→value rows shown under the heading
    body_paragraphs: list[str], # the actual content (timestamps, message text, etc.)
    source_id: str | None = None,  # stable citation token; auto-derived if None
) -> dict:
    """Append one pull section to the current session doc. Creates the doc
    if it doesn't exist yet (lazy-init session = first pull starts it).

    The `source_id` is rendered into the section's metadata block so the
    identification LLM can cite it precisely in each extracted
    requirement. If not supplied, a deterministic ID is derived from
    `source` + `headline` so the same section always gets the same ID."""
    sid = (source_id or _derive_source_id(source, headline)).strip()

    # Cross-process lock around the read-modify-write so concurrent
    # callers (push_chat --all-chats, pull_email, pull_drive, the cron
    # tasks) can't race and corrupt current.docx. The outer pipeline
    # takes file_lock("collect_central"); this is a separate lock name
    # so it never deadlocks against the outer one.
    from tools._lock import file_lock  # lazy to avoid import cycles
    with file_lock("session_doc_append", block=True, wait_max_s=60):
        path = get_or_create()
        doc = Document(str(path))

        # Heading
        h = doc.add_heading(f"{source.upper()} — {headline}", level=1)

        # Metadata block — Source ID first so the LLM can find it easily.
        full_metadata = {"Source ID": sid, **(metadata or {})}
        meta_p = doc.add_paragraph()
        meta_p.paragraph_format.space_after = Pt(2)
        first = True
        for k, v in full_metadata.items():
            if not first:
                meta_p.add_run("   |   ")
            first = False
            kr = meta_p.add_run(f"{k}: ")
            kr.bold = True
            kr.font.size = Pt(9)
            kr.font.color.rgb = _GREY
            vr = meta_p.add_run(str(v))
            vr.font.size = Pt(9)
            vr.font.color.rgb = _GREY

        # Body
        for line in body_paragraphs:
            if not line.strip():
                doc.add_paragraph()
                continue
            doc.add_paragraph(line)

        # Trailing spacer between sections
        doc.add_paragraph()

        doc.save(str(path))

        return {
            "session_path": str(path.relative_to(_NN_ROOT).as_posix()),
            "absolute_path": str(path),
            "section": f"{source.upper()} — {headline}",
            "source_id": sid,
            "appended_paragraphs": len([p for p in body_paragraphs if p.strip()]),
        }


def status() -> dict:
    """Cheap inspection helper — what's currently in the session?"""
    if not SESSION_PATH.exists():
        return {"exists": False}
    meta = _load_meta()
    doc = Document(str(SESSION_PATH))
    section_titles = [p.text for p in doc.paragraphs
                      if p.style.name.startswith("Heading 1")]
    return {
        "exists": True,
        "session_path": str(SESSION_PATH.relative_to(_NN_ROOT).as_posix()),
        "absolute_path": str(SESSION_PATH),
        "started_at": meta.get("started_at"),
        "label": meta.get("label", ""),
        "section_count": len(section_titles),
        "sections": section_titles,
    }
