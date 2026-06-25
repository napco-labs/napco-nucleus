"""
NAPCO Nucleus — Requirement Management tools.

Four MCP tools for the Project Management dimension:

    poll_requirement_emails   IMAP poll from allowlisted senders
    ingest_drive_files        Google Drive → Whisper / pypdf → inbox files
    read_requirement_inbox    Read all raw .txt files Claude will split
    publish_tasks_to_backlog  Create OpenProject Work Packages (dedup
                              against open WPs + SQLite requirements_seen)

Each tool is a thin wrapper around the deterministic module that does
the real work (requirements_inbox / drive_ingester / openproject_client).
Claude-side reasoning happens in the prompts that call these tools.

Memory side-effects (best-effort, never block the primary flow):
    - Every call logs one activity_logs row
    - publish_tasks_to_backlog writes to requirements_seen on BOTH the
      created path AND the dedup-skip-on-already-open path, so a reset
      DB self-heals and updatedRequirements classification can find its
      `updates_prior_id` on subsequent runs.
"""
from __future__ import annotations

import json
import logging
import os
from datetime import datetime, timezone
from pathlib import Path

from claude_agent_sdk import tool

import memory

logger = logging.getLogger(__name__)

_HERE = Path(__file__).parent.parent
_REQ_INBOX_ROOT = _HERE / "data" / "requirements" / "inbox"


def _text(payload) -> dict:
    # ensure_ascii=False so Bangla (call transcripts, client names) reaches
    # the agent as real UTF-8, not \uXXXX escapes. With the json default the
    # identify agent saw escaped Bangla, tried to python/perl-decode it
    # (blocked by the worker sandbox), and bailed with 0 requirements —
    # Assad's two Bangla calls, 2026-06-17.
    return {"content": [{"type": "text",
                         "text": json.dumps(payload, ensure_ascii=False,
                                            default=str)}]}


# ─── poll_requirement_emails ────────────────────────────────────────

@tool(
    "poll_requirement_emails",
    "Fetch new requirement emails from the configured IMAP mailbox and "
    "save each one under data/requirements/inbox/email/ with a short "
    "header (source, from, received, subject, attachment count) "
    "followed by the body. PDF / Word (.docx) / plain text (.txt) "
    "attachments are extracted to text and appended to the body with "
    "`--- attachment: <name> ---` markers, so they read alongside the "
    "message. Filters by REQ_SENDER_ALLOWLIST — only emails from "
    "allowlisted senders are kept. Idempotent via UIDVALIDITY + "
    "since-UID checkpoint. Call this FIRST when the user asks you to "
    "process requirements.",
    {"dry_run": bool},
)
async def poll_requirement_emails_tool(args):
    from mail import requirements_inbox  # lazy — keeps imaplib out of cold-start path
    dry_run = bool(args.get("dry_run", False))
    try:
        result = requirements_inbox.poll_requirement_inbox(dry_run=dry_run)
    except Exception as e:
        logger.exception("poll_requirement_emails failed")
        memory.log_activity(
            task_name="requirement-management:poll_email",
            result=f"error:{type(e).__name__}",
            technical_details={"error": str(e)},
        )
        return _text({"error": str(e), "ingested": 0})

    memory.log_activity(
        task_name="requirement-management:poll_email",
        result=f"ingested:{result.get('ingested', 0)}",
        technical_details={k: v for k, v in result.items() if k != "files"},
    )
    return _text(result)


# ─── ingest_drive_files ─────────────────────────────────────────────

@tool(
    "ingest_drive_files",
    "Download and ingest new files from the configured Google Drive "
    "folder. Audio/video → Groq Whisper → inbox/meetings/. PDF → pypdf "
    "→ inbox/documents/. Word (.docx) → python-docx → inbox/documents/. "
    "Plain text (.txt) → direct read → inbox/documents/. Idempotent — a "
    "Drive file ID is never re-processed. Call this alongside "
    "poll_requirement_emails during the Requirement Management loop so "
    "both sources get captured before read_inbox.",
    {"dry_run": bool},
)
async def ingest_drive_files_tool(args):
    from drive import drive_ingester  # lazy — avoids googleapiclient import on cold start
    dry_run = bool(args.get("dry_run", False))
    try:
        result = drive_ingester.process_new_drive_files(dry_run=dry_run)
    except Exception as e:
        logger.exception("ingest_drive_files failed")
        memory.log_activity(
            task_name="requirement-management:ingest_drive",
            result=f"error:{type(e).__name__}",
            technical_details={"error": str(e)},
        )
        return _text({"error": str(e), "processed": 0})

    memory.log_activity(
        task_name="requirement-management:ingest_drive",
        result=f"processed:{result.get('processed', 0)}",
        technical_details={k: v for k, v in result.items() if k != "files"},
    )
    return _text(result)


# ─── read_requirement_inbox ────────────────────────────────────────

@tool(
    "read_requirement_inbox",
    "Return the contents of every .txt, .md, and .docx file in "
    "data/requirements/inbox/ (email, meetings, chat, documents sub-"
    "folders). Use AFTER the collect-stage tools (poll_requirement_emails, "
    "ingest_drive_files, TRW's pull_chat / transcribe_call) so you can "
    "read every raw requirement source in one call. Email + Drive PDFs "
    "land as .txt; call transcripts land as .md; TRW chat pulls land as "
    ".docx (extracted via python-docx). Optional source= filter: email / "
    "meetings / chat / documents — defaults to all.",
    {"source": str},
)
async def read_requirement_inbox_tool(args):
    source = (args.get("source") or "").strip().lower()
    valid = {"email", "meetings", "chat", "documents"}
    sources = [source] if source in valid else sorted(valid)

    files = []
    for sub in sources:
        sub_dir = _REQ_INBOX_ROOT / sub
        if not sub_dir.is_dir():
            continue
        for name in sorted(os.listdir(sub_dir)):
            if name.startswith("."):
                continue
            lower = name.lower()
            if not lower.endswith((".txt", ".md", ".docx")):
                continue
            path = sub_dir / name
            try:
                if lower.endswith(".docx"):
                    from docx import Document  # lazy
                    doc = Document(str(path))
                    content = "\n\n".join(
                        p.text for p in doc.paragraphs if p.text.strip()
                    )
                else:
                    content = path.read_text(encoding="utf-8")
            except Exception as e:
                logger.warning(f"read {path} failed: {e}")
                continue
            files.append({
                "source": sub,
                "filename": name,
                "rel_path": str(path.relative_to(_HERE)).replace("\\", "/"),
                "chars": len(content),
                "content": content,
            })

    return _text({
        "sources_scanned": sources,
        "file_count": len(files),
        "files": files,
    })


# ─── publish_tasks_to_gitlab ───────────────────────────────────────

def _source_bucket(rel_path: str) -> str:
    """Infer the source bucket from a rel_path for memory recording."""
    p = (rel_path or "").replace("\\", "/")
    for sub in ("meetings", "documents", "chat", "email"):
        if f"/inbox/{sub}/" in p or p.startswith(f"data/requirements/inbox/{sub}/"):
            return sub
    return "email"


@tool(
    "publish_tasks_to_backlog",
    "Create one OpenProject Work Package per task, routed to the project "
    "named by each task's `project` field. Each task must be a dict with keys: "
    "project (REQUIRED — 'mvp-access' or 'cardaccess-4k'; CA4K/MVP aliases "
    "accepted; falls back to the configured default if omitted), title "
    "(required, <70 chars, imperative), description (required — why + "
    "context), acceptance_criteria (optional list of strings, appended "
    "as a bulleted section), estimate_hours (default 3), source_ref "
    "(optional — rel_path of the source file), labels (REQUIRED — "
    "exactly one feature label from {AccessGroup, BadgeHolder, "
    "Personnel} mapped to the WP's Category field, plus exactly one "
    "type label from {requirements, Bug, updatedRequirements}), "
    "issue_type ('task' for requirements/updatedRequirements, 'issue' "
    "for Bug — sets the WP Type), updates_prior_id (REQUIRED for "
    "updatedRequirements, the int Work Package id of the prior WP "
    "being revised). Dedup in two layers: (1) skips any task whose "
    "title already exists as an open WP; (2) skips any task whose "
    "normalized title is already in memory.requirements_seen with a "
    "prior wp_url. Writes to requirements_seen on BOTH paths — newly "
    "created WPs AND exact-title skips — so a reset DB self-heals on "
    "the next run and updatedRequirements classification can find its "
    "prior_id. Honors NAPCO_NUCLEUS_DRY_RUN.",
    {"tasks": list},
)
async def publish_tasks_to_backlog_tool(args):
    import openproject_client as op_client  # lazy

    tasks = args.get("tasks") or []
    if not isinstance(tasks, list) or not tasks:
        return _text({"error": "tasks must be a non-empty list"})

    dry_run = os.environ.get("NAPCO_NUCLEUS_DRY_RUN") == "1"

    # 1. Validate config up front, then pull open Work Packages PER
    # PROJECT lazily. Dedup is project-scoped — a cardaccess-4k title must
    # not dedup against an mvp-access title — so each project gets its own
    # open-title set, fetched on first touch and reused. The skip branch
    # still backfills memory.requirements_seen with id + web_url so a reset
    # DB self-heals and updatedRequirements classification can fire.
    try:
        op_client.default_project()  # raises OpenProjectConfigError if env unset
    except op_client.OpenProjectConfigError as e:
        memory.log_activity(
            task_name="requirement-management:publish_backlog",
            result="error:config_missing",
            technical_details={"error": str(e)},
        )
        return _text({"error": f"OpenProject config missing: {e}",
                      "created": [], "updated": [], "skipped_existing": [], "failed": []})

    # Route each task's free-form `project` to a canonical project slug.
    _PROJECT_ALIASES = {
        "mvp-access": "mvp-access", "mvp access": "mvp-access",
        "mvpaccess": "mvp-access", "mvp": "mvp-access",
        "cardaccess-4k": "cardaccess-4k", "cardaccess 4k": "cardaccess-4k",
        "card-access-4k": "cardaccess-4k", "ca4k": "cardaccess-4k",
        "cardaccess": "cardaccess-4k",
    }

    def _route(task: dict) -> str:
        raw = (task.get("project") or "").strip().lower()
        return _PROJECT_ALIASES.get(raw, raw or op_client.default_project())

    # Per-project open-WP dedup cache: slug -> (open_by_title, open_titles).
    _open_cache: dict[str, tuple[dict, set]] = {}

    def _open_for(project: str) -> tuple[dict, set]:
        if project not in _open_cache:
            wps = op_client.list_open_work_packages(project=project)
            by_title = {i["title"].strip().lower(): i
                        for i in wps if i.get("title")}
            _open_cache[project] = (by_title, set(by_title.keys()))
        return _open_cache[project]

    created, updated, skipped, failed = [], [], [], []

    for t in tasks:
        if not isinstance(t, dict):
            failed.append({"task": t, "error": "not a dict"})
            continue
        title = (t.get("title") or "").strip()
        description = (t.get("description") or "").strip()
        if not title or not description:
            failed.append({"task": t, "error": "title and description required"})
            continue

        # Pull classification fields up front — we need them BEFORE the
        # dedup decision because updatedRequirements goes through a
        # different code path (update existing WP, not create new).
        task_labels = t.get("labels") if isinstance(t.get("labels"), list) else []
        # Accept both new id name and legacy iid name in input dicts.
        is_update = "updatedRequirements" in task_labels
        updates_prior_id = t.get("updates_prior_id") or t.get("updates_prior_iid")
        if not isinstance(updates_prior_id, int):
            updates_prior_id = None

        # Translate label set to OpenProject native fields.
        # Feature label → Category (the project has 3 configured:
        # AccessGroup, BadgeHolder, Personnel).
        _FEATURE_CATEGORIES = {"AccessGroup", "BadgeHolder", "Personnel"}
        op_category = next(
            (l for l in task_labels if l in _FEATURE_CATEGORIES), None,
        )
        # Type label → Work Package Type. Bug rows go to type=Bug; new
        # requirements + revisions both use type=Task. The
        # `updatedRequirements` distinction lives in the comment + the
        # subject change, NOT in a status flip — this OP instance's
        # workflow doesn't permit New → In specification for the
        # service-account role (see openproject_client smoke-test note).
        op_type = "Bug" if "Bug" in task_labels else "Task"

        # Route to the target project and load its open-WP dedup set.
        task_project = _route(t)
        try:
            open_by_title, open_titles = _open_for(task_project)
        except Exception as e:
            logger.exception(f"list_open_work_packages failed for {task_project!r}")
            failed.append({"title": title, "project": task_project,
                           "error": f"could not list project {task_project}: {e}"})
            continue

        # Compose body sections used by both paths.
        body_parts = [description]
        crit = t.get("acceptance_criteria") or []
        if isinstance(crit, list) and crit:
            body_parts.append("\n**Acceptance criteria**")
            for c in crit:
                body_parts.append(f"- {c}")
        est = t.get("estimate_hours") or 3
        body_parts.append(f"\n**Estimate:** ~{est} hours")
        src = (t.get("source_ref") or "").strip()
        if src:
            body_parts.append(f"\n*Source: `{src}`*")
        full_body = "\n".join(body_parts)
        source_bucket = _source_bucket(src)

        # ── UPDATE PATH ─────────────────────────────────────────────
        # updatedRequirements + a known prior id → mutate the existing
        # WP. Title swap + comment with new content. OpenProject's
        # native Activity timeline preserves the full history (subject
        # change is recorded as a journal entry automatically).
        if is_update and updates_prior_id:
            if dry_run:
                updated.append({"id": updates_prior_id,
                                "new_title": title, "dry_run": True})
                memory.remember_requirement(
                    title=title, source=source_bucket, source_ref=src,
                    summary=description[:240],
                    wp_id=updates_prior_id,
                )
                continue
            try:
                op_client.update_work_package(
                    updates_prior_id,
                    title=title,
                    project=task_project,
                    # No status= — workflow rejects New → In spec for
                    # the API user's role on this instance. The subject
                    # change + the timestamped comment below are the
                    # "updated" signal.
                )
                ts = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")
                comment_body = (
                    f"### Update — {ts}\n\n"
                    f"Title changed to: **{title}**\n\n"
                    f"{full_body}"
                )
                op_client.add_work_package_comment(
                    updates_prior_id, comment_body,
                )
                # OP's PATCH response includes _links.self.href; we
                # reconstruct the user-facing URL the same way
                # list_open_work_packages does. But the prior memory
                # row already has the URL, so prefer that.
                prior_url = None
                prior_match = memory.search_requirements(title, limit=1)
                if prior_match:
                    prior_url = prior_match[0].get("wp_url")
                updated.append({"id": updates_prior_id,
                                "new_title": title,
                                "web_url": prior_url})
                memory.remember_requirement(
                    title=title, source=source_bucket, source_ref=src,
                    summary=description[:240],
                    wp_id=updates_prior_id,
                    wp_url=prior_url,
                )
            except Exception as e:
                logger.exception(f"update_work_package failed for id={updates_prior_id}")
                failed.append({"id": updates_prior_id,
                               "title": title, "error": str(e)})
            continue

        # ── CREATE PATH ─────────────────────────────────────────────
        # Standard dedup: skip exact-title matches and fuzzy memory
        # matches with a known id. On exact-title skip, ALSO upsert
        # memory.requirements_seen with the existing WP's id +
        # web_url so a reset DB self-heals — this is what unblocks the
        # updatedRequirements classification on subsequent runs.
        if title.lower() in open_titles:
            existing = open_by_title.get(title.lower(), {})
            memory.remember_requirement(
                title=title, source=source_bucket, source_ref=src,
                summary=description[:240],
                wp_id=existing.get("id"),
                wp_url=existing.get("web_url"),
            )
            skipped.append({
                "title": title,
                "reason": "already open in OpenProject",
                "id": existing.get("id"),
                "web_url": existing.get("web_url"),
                "memory_backfilled": True,
            })
            continue

        prior = memory.search_requirements(title, limit=1)
        if prior and prior[0].get("wp_id"):
            skipped.append({
                "title": title,
                "reason": "already seen in memory",
                "prior_url": prior[0].get("wp_url"),
            })
            continue

        if dry_run:
            created.append({"title": title, "id": "DRY-RUN",
                            "web_url": None, "dry_run": True})
            memory.remember_requirement(
                title=title, source=source_bucket, source_ref=src,
                summary=description[:240],
            )
            open_titles.add(title.lower())
            continue

        try:
            wp = op_client.create_work_package(
                title=title, description=full_body,
                type=op_type, status="New", category=op_category,
                project=task_project,
            )
            wp_id = wp.get("id")
            url = wp.get("web_url")
            created.append({"title": title, "id": wp_id, "web_url": url,
                            "project": task_project})
            memory.remember_requirement(
                title=title, source=source_bucket, source_ref=src,
                summary=description[:240],
                wp_id=wp_id, wp_url=url,
            )
            open_titles.add(title.lower())
        except Exception as e:
            logger.exception(f"create_work_package failed for {title!r}")
            failed.append({"title": title, "error": str(e)})

    memory.log_activity(
        task_name="requirement-management:publish_backlog",
        result=(f"created={len(created)} updated={len(updated)} "
                f"skipped={len(skipped)} failed={len(failed)}"),
        technical_details={
            "created_titles": [c["title"] for c in created],
            "updated_ids":   [u["id"]   for u in updated],
            "skipped_count_by_reason": {
                "open_in_backlog": sum(1 for s in skipped
                                       if s.get("reason") == "already open in OpenProject"),
                "seen_in_memory":  sum(1 for s in skipped
                                       if s.get("reason") == "already seen in memory"),
            },
            "dry_run": dry_run,
        },
    )

    return _text({
        "created": created,
        "updated": updated,
        "skipped_existing": skipped,
        "failed": failed,
        "dry_run": dry_run,
    })


# ─── pull-session helpers ──────────────────────────────────────────

@tool(
    "read_pull_session",
    "Return the contents of the current pull-session Word doc at "
    "data/requirements/sessions/current.docx. Returns the full text "
    "(joined paragraphs) plus a list of section titles so the agent "
    "can see what was pulled. Use this in the verify-session task "
    "instead of read_requirement_inbox — the session doc is the single "
    "input for on-demand identify runs.",
    {},
)
async def read_pull_session_tool(args):
    from tools import _session_doc as session_doc  # lazy
    if not session_doc.SESSION_PATH.exists():
        return _text({"error": "No active pull session — start one with "
                               "start_pull_session and pull data first.",
                      "exists": False})
    from docx import Document  # lazy
    doc = Document(str(session_doc.SESSION_PATH))
    paragraphs = [p.text for p in doc.paragraphs]
    section_titles = [p.text for p in doc.paragraphs
                      if p.style.name.startswith("Heading 1")]
    content = "\n".join(paragraphs)
    meta = session_doc._load_meta()
    return _text({
        "exists": True,
        "session_path": str(session_doc.SESSION_PATH.relative_to(_HERE)
                            .as_posix()),
        "started_at": meta.get("started_at"),
        "label": meta.get("label", ""),
        "section_count": len(section_titles),
        "sections": section_titles,
        "chars": len(content),
        "content": content,
    })


@tool(
    "start_pull_session",
    "Archive the current pull-session doc (if any) and start a fresh one. "
    "Optional `label` argument is stored in meta and used in the archive "
    "filename. Use when the user says 'start a new session'.",
    {"label": str},
)
async def start_pull_session_tool(args):
    from tools import _session_doc as session_doc  # lazy
    label = (args.get("label") or "").strip() or None
    result = session_doc.reset(label=label)
    memory.log_activity(
        task_name="requirement-collection:start_pull_session",
        result="reset",
        technical_details=result,
    )
    return _text(result)


@tool(
    "pull_session_status",
    "Inspect the current pull-session doc — returns whether it exists, "
    "when it was started, label (if any), and the list of section "
    "titles already in it. No side effects.",
    {},
)
async def pull_session_status_tool(args):
    from tools import _session_doc as session_doc  # lazy
    return _text(session_doc.status())


TOOLS = [
    poll_requirement_emails_tool,
    ingest_drive_files_tool,
    read_requirement_inbox_tool,
    publish_tasks_to_backlog_tool,
    read_pull_session_tool,
    start_pull_session_tool,
    pull_session_status_tool,
]

TOOL_NAMES = [
    "poll_requirement_emails",
    "ingest_drive_files",
    "read_requirement_inbox",
    "publish_tasks_to_backlog",
    "read_pull_session",
    "start_pull_session",
    "pull_session_status",
]
